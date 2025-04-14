import os
import re
from datetime import datetime
from docx import Document

# Output folder path
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Helper to clean filenames
def sanitize_filename(name: str) -> str:
    return re.sub(r'[^a-zA-Z0-9_\-]', '', name.replace(" ", "_"))

# Enhanced Markdown-to-DOCX parser with table support
def parse_markdown_to_docx(doc: Document, content: str):
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Horizontal rule (---)
        if re.match(r"^-{3,}$", line):
            doc.add_paragraph("")
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            quote_line = re.sub(r"^> ", "", line)
            quote_line = re.sub(r"\*\*(.*?)\*\*", r"\1", quote_line)
            quote_line = re.sub(r"\*(.*?)\*", r"\1", quote_line)
            doc.add_paragraph(quote_line, style='Intense Quote')
            i += 1
            continue

        # Table block (starts with "|")
        if line.startswith("|") and "|" in line:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            # Remove separator row (like |---|---|)
            if len(table_lines) >= 2 and re.match(r"^\s*\|[-| ]+\|\s*$", table_lines[1]):
                table_lines.pop(1)

            # Parse table
            rows = [re.split(r'\s*\|\s*', row.strip('| ')) for row in table_lines]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    cell_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell)
                    table.cell(r_idx, c_idx).text = cell_text
            continue

        # Bold/Italic cleanup
        clean_line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        clean_line = re.sub(r"\*(.*?)\*", r"\1", clean_line)

        # Headers
        header_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if header_match:
            level = len(header_match.group(1))
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", header_match.group(2))
            doc.add_heading(text, level=min(level, 4))
        elif line.startswith("- "):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        else:
            doc.add_paragraph(clean_line)

        i += 1

# Main output agent function
def run_output_agent(final_article: str, draft_title: str) -> dict:
    print("📝 OutputAgent received article and title...")
    print("📌 Draft Title:", draft_title)

    # Clean filename
    safe_title = sanitize_filename(draft_title)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    # Create and format Word doc
    doc = Document()
    parse_markdown_to_docx(doc, final_article)
    doc.save(filepath)

    # HTML preview (optional)
    html_preview = final_article.replace("\n", "<br>")

    return {
        "agent": "OutputAgent",
        "markdown": final_article,
        "docx_file_path": filepath,
        "html": html_preview
    }
